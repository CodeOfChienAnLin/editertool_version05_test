import os
import sys
import json
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import threading
import docx2txt  # 用於讀取Word文檔
import msoffcrypto  # 用於處理加密的Office文檔
import io
from io import BytesIO
import opencc  # 用於中文文字轉換和校正
import tempfile
from docx import Document  # 用於更精確地讀取Word文檔格式

class TextCorrectionTool:
    """文字校正工具主類別"""
    def __init__(self, root):
        """初始化應用程式
        
        參數:
            root: tkinter的根視窗
        """
        self.root = root
        self.root.title("文字校正工具")
        self.root.geometry("900x600")  # 設定視窗大小為900x600
        self.root.resizable(False, False)  # 禁止調整視窗大小
        
        # 載入詞彙保護表
        self.protected_words = self.load_protected_words()
        
        # 載入設定
        self.settings = self.load_settings()
        
        # 初始化OpenCC轉換器
        try:
            # 使用簡體到繁體的轉換
            self.converter = opencc.OpenCC('s2t')  # 將簡體字轉為繁體字
        except Exception as e:
            messagebox.showerror("錯誤", f"無法初始化OpenCC轉換器: {str(e)}")
            self.converter = None
        
        self.create_widgets()  # 創建UI元件
        self.setup_drag_drop()  # 設置拖放功能
    
    def create_widgets(self):
        """創建所有UI元件"""
        # 選單列
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 檔案選單
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="檔案", menu=file_menu)
        file_menu.add_command(label="開啟", command=self.open_file)
        file_menu.add_command(label="儲存", command=self.save_file)
        file_menu.add_separator()
        file_menu.add_command(label="離開", command=self.root.quit)
        
        # 編輯選單
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="編輯", menu=edit_menu)
        edit_menu.add_command(label="校正文字", command=self.correct_text)
        edit_menu.add_command(label="管理保護詞彙", command=self.manage_protected_words)
        
        # 設定選單
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="設定", menu=settings_menu)
        settings_menu.add_command(label="文字格式", command=self.open_text_settings)
        
        # 文字處理區域框架 (700x500)
        text_frame = tk.Frame(self.root, width=700, height=500)
        text_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 添加垂直滾動條
        y_scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 文字處理區域 - 用於顯示和編輯文字
        self.text_area = tk.Text(text_frame, 
                               font=(self.settings["font_family"], self.settings["font_size"]),
                               wrap=tk.WORD,  # 啟用自動換行
                               yscrollcommand=y_scrollbar.set)
        self.text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 設置縮進，使換行後的文字對齊前一行的第一個字
        self.text_area.config(tabs=("1c", "2c", "3c", "4c"), tabstyle="wordprocessor")
        
        # 綁定事件，當文字變化時調整縮進
        self.text_area.bind("<<Modified>>", self.adjust_indentation)
        
        # 設置滾動條的命令
        y_scrollbar.config(command=self.text_area.yview)
        
        # 圖片顯示區域 (900x100)
        image_frame = tk.Frame(self.root, width=900, height=100, bg="lightgrey")
        image_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 添加標籤到圖片區域
        self.image_label = tk.Label(image_frame, text="圖片顯示區域", bg="lightgrey")
        self.image_label.pack(fill=tk.BOTH, expand=True)
        
        # 狀態列
        self.status_bar = tk.Label(self.root, text="就緒", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def setup_drag_drop(self):
        """設置拖放功能"""
        try:
            # 直接使用 Tkinter 原生的拖放功能
            # 為文字區域啟用拖放
            self.text_area.drop_target_register('DND_Files')
            self.text_area.dnd_bind('<<Drop>>', self.handle_drop)
            self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
            print("已啟用 Tkinter 原生拖放功能")
        except Exception as e:
            print(f"Tkinter 原生拖放初始化失敗: {str(e)}")
            
            # 嘗試使用 TkDND
            try:
                print("嘗試使用 TkDND...")
                # 嘗試將 TkDND 套件目錄加入路徑
                import sys
                import os
                tkdnd_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), 'tkdnd'))
                if os.path.exists(tkdnd_dir):
                    sys.path.append(tkdnd_dir)
                
                from tkinter import TkVersion
                if TkVersion >= 8.6:
                    # Tk 8.6+ 原生支援拖放
                    self.root.tk.call('package', 'require', 'tkdnd')
                    self.root.tk.call('tkdnd::drop_target', 'register', self.text_area._w)
                    self.root.tk.call('bind', self.text_area._w, '<<Drop>>', 
                                     self.root.register(self.handle_drop))
                    print("使用 Tk 8.6+ 原生拖放功能")
                    self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                    return
                
                # 嘗試使用 TkinterDnD2
                try:
                    print("嘗試使用 TkinterDnD2...")
                    # 使用絕對導入
                    from TkinterDnD2 import TkinterDnD, DND_FILES
                    TkinterDnD.dnd_start(self.root)
                    self.text_area.drop_target_register(DND_FILES)
                    self.text_area.dnd_bind('<<Drop>>', self.handle_drop)
                    print("使用 TkinterDnD2 拖放功能")
                    self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                    return
                except Exception as e:
                    print(f"TkinterDnD2 初始化失敗: {str(e)}")
                
                # 嘗試使用自訂的 TkDND 包裝類
                try:
                    print("嘗試使用自訂 TkDND 包裝類...")
                    from tkdnd_wrapper import TkDND
                    dnd = TkDND(self.root)
                    success = dnd.bindtarget(self.text_area, self.handle_drop, 'text/uri-list')
                    if success:
                        print("使用自訂 TkDND 包裝類")
                        self.status_bar.config(text="拖放功能已啟用，可以拖放 Word 文檔")
                        return
                except Exception as e:
                    print(f"自訂 TkDND 包裝類初始化失敗: {str(e)}")
                
            except Exception as e:
                print(f"TkDND 相關初始化失敗: {str(e)}")
            
            # 最後嘗試使用簡易的捕獲方法
            try:
                print("嘗試使用簡易捕獲方法...")
                # 處理貼上事件
                self.root.bind("<FocusIn>", self.check_clipboard)
                self.root.bind("<ButtonRelease>", self.check_clipboard)
                self.root.bind("<Key>", lambda e: self.check_clipboard() if e.keysym == 'v' and (e.state & 4) else None)
                print("已啟用簡易捕獲方法")
                self.status_bar.config(text="已啟用替代拖放功能，將檔案拖放到視窗後請點擊")
                return
            except Exception as e:
                print(f"簡易捕獲方法初始化失敗: {str(e)}")
            
            # 所有方法都失敗
            print("所有拖放方法都失敗了")
            self.status_bar.config(text="拖放功能初始化失敗，請使用選單開啟檔案")
            messagebox.showwarning("拖放功能警告", "拖放功能無法初始化\n請使用選單開啟檔案")
    
    def check_clipboard(self, event=None):
        """檢查剪貼簿是否有檔案路徑"""
        try:
            clipboard = self.root.clipboard_get()
            if clipboard and os.path.exists(clipboard) and clipboard.lower().endswith(('.docx', '.doc')):
                print(f"從剪貼簿獲取檔案: {clipboard}")
                self.process_word_file(clipboard)
                return True
        except Exception as e:
            print(f"檢查剪貼簿時發生錯誤: {str(e)}")
        return False
    
    def handle_drop(self, event):
        """處理檔案拖放事件
        
        參數:
            event: 拖放事件物件
        """
        try:
            data = event.data
            file_path = str(data).strip()
            
            print(f"原始拖放路徑: {file_path}")
            
            # 處理可能的格式
            # Windows 可能會在路徑周圍添加大括號或引號
            if file_path.startswith('{') and file_path.endswith('}'):
                file_path = file_path[1:-1]
            
            # 移除可能的引號
            if (file_path.startswith('"') and file_path.endswith('"')) or \
               (file_path.startswith("'") and file_path.endswith("'")):
                file_path = file_path[1:-1]
            
            # 處理可能的檔案URL格式
            if file_path.startswith('file:///'):
                file_path = file_path[8:].replace('/', '\\')
            
            # 處理 Mac 路徑格式或其他非標準路徑
            if file_path.startswith('/Mac/') or '://' in file_path:
                # 嘗試從路徑中提取實際的文件名
                file_name = os.path.basename(file_path)
                
                # 顯示錯誤訊息
                messagebox.showinfo("路徑格式不支援", 
                                   f"檢測到非標準路徑格式: {file_path}\n\n"
                                   f"請嘗試以下方法：\n"
                                   f"1. 使用「檔案」選單中的「開啟」功能\n"
                                   f"2. 從檔案總管直接拖放檔案\n"
                                   f"3. 確保檔案位於本機上，而非網路位置")
                return
            
            print(f"處理後的檔案路徑: {file_path}")
            
            # 檢查檔案是否存在
            if not os.path.exists(file_path):
                messagebox.showerror("錯誤", f"找不到檔案: {file_path}\n請確保檔案路徑正確且檔案存在。")
                return
                
            # 檢查檔案是否為Word檔案
            if not file_path.lower().endswith(('.doc', '.docx')):
                messagebox.showerror("錯誤", f"不支援的檔案格式: {file_path}\n僅支援 .doc 和 .docx 格式。")
                return
                
            # 更新狀態欄
            self.status_bar.config(text=f"正在處理檔案: {os.path.basename(file_path)}")
            
            # 嘗試處理Word檔案
            try:
                # 先嘗試檢查文件是否加密
                try:
                    with open(file_path, 'rb') as f:
                        try:
                            office_file = msoffcrypto.OfficeFile(f)
                            if office_file.is_encrypted():
                                print("檔案已加密，需要密碼")
                                # 文件已加密，直接調用密碼處理方法
                                self.handle_password_protected_file(file_path)
                                return
                        except Exception as e:
                            print(f"檢查加密狀態時發生錯誤: {str(e)}")
                            # 繼續嘗試普通處理
                except Exception as e:
                    print(f"開啟檔案時發生錯誤: {str(e)}")
                    # 繼續嘗試普通處理
                
                # 嘗試不使用密碼處理
                text = self.process_word_file(file_path)
                
                # 如果成功處理，更新文字區域
                if text:
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, text)
                    self.status_bar.config(text=f"已載入檔案: {os.path.basename(file_path)}")
                    
                    # 調整縮進
                    self.adjust_indentation()
                    
                    # 自動校正文字
                    self.correct_text()
                
            except Exception as e:
                # 檢查是否為加密文件的錯誤
                error_str = str(e).lower()
                if self._is_password_error(error_str):
                    # 可能是加密文件，嘗試使用密碼處理
                    print(f"檢測到加密錯誤: {error_str}")
                    self.handle_password_protected_file(file_path)
                else:
                    # 其他錯誤，顯示錯誤訊息
                    messagebox.showerror("錯誤", f"處理檔案時發生錯誤: {str(e)}")
                    self.status_bar.config(text=f"處理檔案時發生錯誤: {str(e)}")
                
        except Exception as e:
            print(f"處理拖放檔案時發生錯誤: {str(e)}")
            self.status_bar.config(text=f"處理拖放檔案時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"處理拖放檔案時發生錯誤: {str(e)}")
    
    def process_word_file(self, file_path, password=None):
        """處理Word檔案
        
        參數:
            file_path: Word檔案路徑
            password: 檔案密碼（如果有的話）
        
        回傳:
            檔案內容
        """
        try:
            # 檢查檔案是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"找不到檔案: {file_path}")
            
            # 檢查檔案是否為Word檔案
            if not file_path.lower().endswith(('.doc', '.docx')):
                raise ValueError(f"不支援的檔案格式: {file_path}")
            
            # 更新狀態欄
            self.status_bar.config(text=f"正在處理檔案: {os.path.basename(file_path)}")
            
            # 如果有密碼，嘗試解密
            if password:
                try:
                    with open(file_path, 'rb') as encrypted_file:
                        # 使用 msoffcrypto 解密
                        office_file = msoffcrypto.OfficeFile(encrypted_file)
                        
                        # 檢查文件是否加密
                        if not office_file.is_encrypted():
                            print("文件未加密，無需解密")
                            # 如果文件未加密，直接處理
                            return self._process_unencrypted_file(file_path)
                        
                        # 解密文件到內存
                        decrypted_content = BytesIO()
                        try:
                            office_file.load_key(password=password)
                            office_file.decrypt(decrypted_content)
                        except Exception as e:
                            print(f"解密失敗: {str(e)}")
                            raise ValueError(f"解密失敗，密碼可能不正確: {str(e)}")
                        
                        # 重置指針到開始位置
                        decrypted_content.seek(0)
                        
                        # 嘗試使用 python-docx 解析解密後的內容
                        try:
                            doc = Document(decrypted_content)
                            text = self._extract_text_from_document(doc)
                            return text
                        except Exception as docx_e:
                            print(f"使用 python-docx 解析失敗: {str(docx_e)}")
                            
                            # 如果 python-docx 失敗，嘗試使用臨時文件和 docx2txt
                            try:
                                # 創建臨時文件
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                                temp_file.close()
                                
                                # 將解密內容寫入臨時文件
                                decrypted_content.seek(0)
                                with open(temp_file.name, 'wb') as f:
                                    f.write(decrypted_content.read())
                                
                                # 使用 docx2txt 處理
                                try:
                                    text = docx2txt.process(temp_file.name)
                                    return text
                                finally:
                                    # 刪除臨時文件
                                    if os.path.exists(temp_file.name):
                                        os.unlink(temp_file.name)
                            except Exception as temp_e:
                                print(f"使用臨時文件處理失敗: {str(temp_e)}")
                                raise ValueError(f"無法解析解密後的文件: {str(temp_e)}")
                except Exception as e:
                    print(f"處理加密文件時發生錯誤: {str(e)}")
                    raise ValueError(f"處理加密文件時發生錯誤: {str(e)}")
            else:
                # 處理未加密文件
                return self._process_unencrypted_file(file_path)
                
        except Exception as e:
            print(f"處理檔案時發生錯誤: {str(e)}")
            self.status_bar.config(text=f"處理檔案時發生錯誤: {str(e)}")
            
            # 檢查是否為加密錯誤
            if password is None and self._is_password_error(str(e)):
                raise Exception(f"檔案可能有密碼保護: {str(e)}")
            
            raise e
    
    def _process_unencrypted_file(self, file_path):
        """處理未加密的Word檔案
        
        參數:
            file_path: Word檔案路徑
            
        回傳:
            檔案內容
        """
        # 先嘗試使用 docx2txt
        try:
            text = docx2txt.process(file_path)
            if text:
                return text
        except Exception as e:
            print(f"使用 docx2txt 處理失敗: {str(e)}")
            
            # 如果是加密錯誤，直接拋出
            if self._is_password_error(str(e)):
                raise Exception(f"檔案可能有密碼保護: {str(e)}")
        
        # 如果 docx2txt 失敗，嘗試使用 python-docx
        try:
            doc = Document(file_path)
            text = self._extract_text_from_document(doc)
            if text:
                return text
        except Exception as docx_e:
            print(f"使用 python-docx 處理失敗: {str(docx_e)}")
            
            # 如果是加密錯誤，直接拋出
            if self._is_password_error(str(docx_e)):
                raise Exception(f"檔案可能有密碼保護: {str(docx_e)}")
            
            # 如果兩種方法都失敗，則拋出異常
            raise Exception(f"無法讀取文件: {str(docx_e)}")
    
    def _is_password_error(self, error_message):
        """檢查錯誤訊息是否與密碼保護相關
        
        參數:
            error_message: 錯誤訊息
            
        回傳:
            是否為密碼相關錯誤
        """
        error_message = error_message.lower()
        password_keywords = ["password", "encrypted", "保護", "密碼", "加密"]
        return any(keyword in error_message for keyword in password_keywords)
    
    def _extract_text_from_document(self, doc):
        """從 python-docx Document 物件中提取文字
        
        參數:
            doc: python-docx Document 物件
            
        回傳:
            提取的文字
        """
        # 提取文本，保留段落格式
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():  # 忽略空段落
                paragraphs.append(para.text)
        
        # 提取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append('\t'.join(row_text))
        
        # 使用兩個換行符連接段落，保留格式
        return '\n\n'.join(paragraphs)
    
    def handle_password_protected_file(self, file_path):
        """處理有密碼保護的Word檔案
        
        參數:
            file_path: 加密Word檔案的路徑
        """
        # 處理有密碼保護的檔案
        password = self.ask_password()
        if password:
            try:
                # 使用密碼解密檔案
                text = self.process_word_file(file_path, password)
                self.text_area.delete(1.0, tk.END)
                self.text_area.insert(tk.END, text)
                self.status_bar.config(text=f"已載入加密檔案: {os.path.basename(file_path)}")
                
                # 調整縮進
                self.adjust_indentation()
                
                # 自動校正文字
                self.correct_text()
            except Exception as e:
                messagebox.showerror("錯誤", f"解密失敗，密碼可能不正確: {str(e)}")
                self.status_bar.config(text=f"解密失敗: {os.path.basename(file_path)}")
    
    def ask_password(self):
        """顯示密碼輸入對話框
        
        回傳:
            使用者輸入的密碼
        """
        # 創建密碼輸入對話框
        password_window = tk.Toplevel(self.root)
        password_window.title("密碼保護")
        password_window.geometry("300x150")
        password_window.resizable(False, False)
        
        # 設置模態對話框
        password_window.transient(self.root)
        password_window.grab_set()
        
        # 居中顯示
        window_width = 300
        window_height = 150
        screen_width = password_window.winfo_screenwidth()
        screen_height = password_window.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        password_window.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # 添加說明標籤
        tk.Label(password_window, text="該檔案有密碼保護，請輸入密碼:", font=("Arial", 10)).pack(pady=10)
        
        # 密碼輸入框
        password_entry = tk.Entry(password_window, show="*", width=25)
        password_entry.pack(pady=5)
        password_entry.focus_set()  # 設置焦點
        
        password = None
        
        # 確定按鈕回調函數
        def on_ok():
            nonlocal password
            password = password_entry.get()
            password_window.destroy()
        
        # 取消按鈕回調函數
        def on_cancel():
            password_window.destroy()
        
        # 按鈕區域
        button_frame = tk.Frame(password_window)
        button_frame.pack(pady=10)
        
        tk.Button(button_frame, text="確定", command=on_ok, width=10).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="取消", command=on_cancel, width=10).pack(side=tk.LEFT, padx=5)
        
        # 綁定回車鍵
        password_window.bind("<Return>", lambda event: on_ok())
        password_window.bind("<Escape>", lambda event: on_cancel())
        
        # 等待視窗關閉
        password_window.wait_window()
        return password
    
    def open_file(self):
        """開啟檔案對話框"""
        try:
            file_path = filedialog.askopenfilename(
                title="選擇Word檔案", 
                filetypes=[("Word文件", "*.docx;*.doc"), ("所有檔案", "*.*")]
            )
            
            if file_path:
                print(f"選擇的檔案: {file_path}")
                self.status_bar.config(text=f"選擇的檔案: {file_path}")
                text = self.process_word_file(file_path)
                if text:
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, text)
                    self.status_bar.config(text=f"已載入檔案: {os.path.basename(file_path)}")
                
                # 調整縮進
                self.adjust_indentation()
                
                # 自動校正文字
                self.correct_text()
                
        except Exception as e:
            print(f"開啟檔案錯誤: {str(e)}")
            self.status_bar.config(text="開啟檔案時出錯")
            messagebox.showerror("錯誤", f"無法開啟檔案: {str(e)}")
    
    def save_file(self):
        """儲存檔案對話框"""
        # 儲存檔案對話框
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("文字檔案", "*.txt"), ("Word文檔", "*.docx"), ("所有檔案", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    text = self.text_area.get(1.0, tk.END)
                    f.write(text)
                self.status_bar.config(text=f"已儲存到: {os.path.basename(file_path)}")
            except Exception as e:
                self.status_bar.config(text="儲存檔案時出錯")
                messagebox.showerror("錯誤", f"無法儲存檔案: {str(e)}")
    
    def correct_text(self):
        """校正文字內容"""
        # 檢查OpenCC轉換器是否正確初始化
        if not self.converter:
            self.status_bar.config(text="OpenCC轉換器未正確初始化，無法進行校正")
            messagebox.showerror("錯誤", "OpenCC轉換器未正確初始化，無法進行校正")
            return
            
        # 獲取文字內容
        text = self.text_area.get(1.0, tk.END)
        
        # 在背景執行校正，避免UI凍結
        threading.Thread(target=self._correct_text_thread, args=(text,)).start()
    
    def _correct_text_thread(self, text):
        """在背景執行文字校正的執行緒
        
        參數:
            text: 要校正的文字
        """
        try:
            print("開始文字校正執行緒")
            
            # 載入保護詞彙
            protected_words = self.load_protected_words()
            print(f"已載入保護詞彙: {protected_words}")
            
            # 創建一個暫時字典來保存保護詞彙的位置
            protected_positions = {}
            
            # 找出所有保護詞彙在文本中的位置
            for word in protected_words:
                start_pos = 0
                while True:
                    pos = text.find(word, start_pos)
                    if pos == -1:
                        break
                    protected_positions[pos] = (pos + len(word), word)
                    start_pos = pos + 1
            
            print(f"找到 {len(protected_positions)} 個保護詞彙位置")
            
            # 如果沒有保護詞彙，直接轉換整個文本
            if not protected_positions:
                corrected_text = self.converter.convert(text)
            else:
                # 分段處理文本，保護特定詞彙
                result = []
                last_end = 0
                
                # 按位置排序保護區域
                positions = sorted(protected_positions.keys())
                
                for start in positions:
                    end, word = protected_positions[start]
                    
                    # 轉換保護詞彙前的文本
                    if start > last_end:
                        segment = text[last_end:start]
                        result.append(self.converter.convert(segment))
                    
                    # 添加保護詞彙（不轉換）
                    result.append(word)
                    last_end = end
                
                # 處理最後一個保護詞彙之後的文本
                if last_end < len(text):
                    segment = text[last_end:]
                    result.append(self.converter.convert(segment))
                
                corrected_text = ''.join(result)
            
            print(f"校正完成，轉換後文字長度: {len(corrected_text)}")
            
            # 更新UI必須在主執行緒中進行
            self.root.after(0, self._update_text_area, corrected_text)
        except Exception as e:
            print(f"校正文字時發生錯誤: {str(e)}")
            # 更新UI必須在主執行緒中進行
            self.root.after(0, lambda: self.status_bar.config(text=f"校正文字時發生錯誤: {str(e)}"))
            self.root.after(0, lambda: messagebox.showerror("錯誤", f"校正文字時發生錯誤: {str(e)}"))
    
    def _update_text_area(self, corrected_text):
        """更新文字區域的內容
        
        參數:
            corrected_text: 校正後的文字
        """
        self.text_area.delete(1.0, tk.END)
        self.text_area.insert(tk.END, corrected_text)
        self.status_bar.config(text="文字校正完成")
    
    def load_protected_words(self):
        """載入詞彙保護表
        
        回傳:
            詞彙保護列表
        """
        try:
            # 檢查檔案是否存在
            if not os.path.exists("protected_words.json"):
                # 如果不存在，創建一個空的詞彙保護表
                with open("protected_words.json", "w", encoding="utf-8") as f:
                    json.dump({"protected_words": []}, f, ensure_ascii=False, indent=4)
                return []
            
            # 讀取詞彙保護表
            with open("protected_words.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                
            # 確保 protected_words 鍵存在
            if "protected_words" not in data:
                return []
                
            return data["protected_words"]
        except Exception as e:
            print(f"載入詞彙保護表時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法載入詞彙保護表: {str(e)}")
            return []
    
    def save_protected_words(self):
        """儲存詞彙保護表"""
        # 儲存詞彙保護表
        try:
            with open('protected_words.json', 'w', encoding='utf-8') as f:
                json.dump(self.protected_words, f, ensure_ascii=False, indent=4)
        except Exception as e:
            messagebox.showerror("錯誤", f"無法儲存詞彙保護表: {str(e)}")
    
    def manage_protected_words(self):
        """管理保護詞彙的視窗"""
        # 創建一個新視窗來管理保護詞彙
        manage_window = tk.Toplevel(self.root)
        manage_window.title("管理保護詞彙")
        manage_window.geometry("400x500")
        
        # 創建一個框架
        frame = tk.Frame(manage_window)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 添加標籤
        tk.Label(frame, text="保護詞彙列表:").pack(anchor=tk.W)
        
        # 添加列表框和滾動條
        list_frame = tk.Frame(frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        words_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set)
        words_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=words_listbox.yview)
        
        # 填充列表框
        for word in self.protected_words:
            words_listbox.insert(tk.END, word)
        
        # 添加輸入欄位和按鈕
        input_frame = tk.Frame(frame)
        input_frame.pack(fill=tk.X, pady=5)
        
        tk.Label(input_frame, text="新增詞彙:").pack(side=tk.LEFT)
        word_entry = tk.Entry(input_frame)
        word_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # 添加按鈕
        buttons_frame = tk.Frame(frame)
        buttons_frame.pack(fill=tk.X)
        
        def add_word():
            """添加新詞彙到保護列表"""
            word = word_entry.get().strip()
            if word and word not in self.protected_words:
                self.protected_words.append(word)
                words_listbox.insert(tk.END, word)
                word_entry.delete(0, tk.END)
                self.save_protected_words()
        
        def remove_word():
            """從保護列表中移除選中的詞彙"""
            selection = words_listbox.curselection()
            if selection:
                index = selection[0]
                word = words_listbox.get(index)
                words_listbox.delete(index)
                self.protected_words.remove(word)
                self.save_protected_words()
        
        tk.Button(buttons_frame, text="添加", command=add_word).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="刪除", command=remove_word).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="關閉", command=manage_window.destroy).pack(side=tk.RIGHT, padx=5)

    def open_text_settings(self):
        """開啟文字格式設定視窗"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("文字格式設定")
        settings_window.geometry("400x300")
        settings_window.resizable(False, False)
        settings_window.transient(self.root)  # 設為主視窗的子視窗
        settings_window.grab_set()  # 模態視窗
        
        # 建立框架
        frame = tk.Frame(settings_window, padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)
        
        # 字體選擇
        tk.Label(frame, text="字體:").grid(row=0, column=0, sticky=tk.W, pady=10)
        
        # 獲取系統可用字體
        available_fonts = ["新細明體", "標楷體", "微軟正黑體", "Arial", "Times New Roman", "Courier New"]
        
        font_var = tk.StringVar(value=self.settings["font_family"])
        font_combo = ttk.Combobox(frame, textvariable=font_var, values=available_fonts, width=20)
        font_combo.grid(row=0, column=1, sticky=tk.W, pady=10)
        
        # 字體大小選擇
        tk.Label(frame, text="字體大小:").grid(row=1, column=0, sticky=tk.W, pady=10)
        
        size_var = tk.IntVar(value=self.settings["font_size"])
        size_combo = ttk.Combobox(frame, textvariable=size_var, values=[8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 26, 28, 36], width=5)
        size_combo.grid(row=1, column=1, sticky=tk.W, pady=10)
        
        # 預覽區域
        tk.Label(frame, text="預覽:").grid(row=2, column=0, sticky=tk.W, pady=10)
        
        preview_text = tk.Text(frame, width=30, height=5, wrap=tk.WORD)
        preview_text.grid(row=2, column=1, sticky=tk.W, pady=10)
        preview_text.insert(tk.END, "這是預覽文字\nABCDEFG\n123456789")
        
        # 更新預覽的函數
        def update_preview(*args):
            font_family = font_var.get()
            font_size = size_var.get()
            preview_text.configure(font=(font_family, font_size))
        
        # 綁定變更事件
        font_var.trace_add("write", update_preview)
        size_var.trace_add("write", update_preview)
        
        # 初始更新預覽
        update_preview()
        
        # 按鈕區域
        button_frame = tk.Frame(frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=20)
        
        # 確定按鈕
        def save_settings():
            self.settings["font_family"] = font_var.get()
            self.settings["font_size"] = size_var.get()
            self.save_settings()
            self.text_area.configure(font=(self.settings["font_family"], self.settings["font_size"]))
            settings_window.destroy()
            
        tk.Button(button_frame, text="確定", command=save_settings, width=10).pack(side=tk.LEFT, padx=10)
        
        # 取消按鈕
        tk.Button(button_frame, text="取消", command=settings_window.destroy, width=10).pack(side=tk.LEFT, padx=10)
    
    def load_settings(self):
        """載入設定
        
        回傳:
            設定字典
        """
        default_settings = {
            "font_family": "新細明體",
            "font_size": 12
        }
        
        try:
            # 檢查檔案是否存在
            if not os.path.exists("settings.json"):
                # 如果不存在，創建一個預設設定檔
                with open("settings.json", "w", encoding="utf-8") as f:
                    json.dump(default_settings, f, ensure_ascii=False, indent=4)
                return default_settings
            
            # 讀取設定檔
            with open("settings.json", "r", encoding="utf-8") as f:
                settings = json.load(f)
                
            # 確保所有必要的設定都存在
            for key in default_settings:
                if key not in settings:
                    settings[key] = default_settings[key]
                    
            return settings
        except Exception as e:
            print(f"載入設定時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法載入設定: {str(e)}")
            return default_settings
    
    def save_settings(self):
        """儲存設定"""
        try:
            with open("settings.json", "w", encoding="utf-8") as f:
                json.dump(self.settings, f, ensure_ascii=False, indent=4)
            print("設定已儲存")
        except Exception as e:
            print(f"儲存設定時發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"無法儲存設定: {str(e)}")

    def adjust_indentation(self, event=None):
        """調整文字縮進，使換行後的文字對齊前一行的第一個字"""
        # 重置修改標誌，避免無限循環
        self.text_area.edit_modified(False)
        
        # 獲取所有文字
        content = self.text_area.get("1.0", tk.END)
        
        # 如果內容為空，不做處理
        if not content.strip():
            return
        
        # 處理每個段落
        lines = content.split('\n')
        for i in range(len(lines)):
            # 跳過空行
            if not lines[i].strip():
                continue
                
            # 獲取當前行第一個非空白字符的位置
            first_char_pos = len(lines[i]) - len(lines[i].lstrip())
            
            # 如果不是第一行且前一行不為空，設置縮進
            if i > 0 and lines[i-1].strip():
                prev_first_char_pos = len(lines[i-1]) - len(lines[i-1].lstrip())
                
                # 如果當前行是前一行的換行部分（由自動換行產生）
                # 這裡需要根據實際情況調整判斷邏輯
                if first_char_pos == 0 and len(lines[i]) > 0:
                    # 設置縮進標籤
                    tag_name = f"indent_{i}"
                    self.text_area.tag_configure(tag_name, lmargin1=prev_first_char_pos)
                    
                    # 應用標籤到當前行
                    line_start = f"{i+1}.0"
                    line_end = f"{i+1}.{len(lines[i])}"
                    self.text_area.tag_add(tag_name, line_start, line_end)

    def adjust_text_formatting(self, event=None):
        """調整文字格式，包括縮進和對齊"""
        # 調用原有的縮進方法
        self.adjust_indentation(event)


def main():
    """程式主入口點"""
    try:
        # 嘗試使用 TkinterDnD2 創建支援拖放的根視窗
        try:
            from tkinterdnd2 import TkinterDnD, DND_FILES
            root = TkinterDnD.Tk()
            print("成功使用 TkinterDnD2 初始化根視窗")
        except Exception as e:
            print(f"無法使用 TkinterDnD2: {str(e)}")
            # 退回使用普通的 Tk
            root = tk.Tk()
            print("使用普通 Tk 初始化根視窗")
        
        app = TextCorrectionTool(root)
        root.mainloop()
    except Exception as e:
        print(f"程式執行錯誤: {str(e)}")
        messagebox.showerror("錯誤", f"程式執行錯誤: {str(e)}")

if __name__ == "__main__":
    main()
